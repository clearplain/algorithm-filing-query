from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import urljoin

import requests
from bs4 import BeautifulSoup
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOMESTIC_DIR = ROOT / "境内互联网信息服务算法备案清单（已下载）"
DEEP_DIR = ROOT / "深度合成算法备案信息（未下载）"
OUT_DIR = ROOT / "public" / "data"

DOMESTIC_SOURCE = "https://www.cac.gov.cn/2022-08/12/c_1661927474338504.htm"
ALGO_LAW_SOURCE = "https://www.12377.cn/bmgz/2022/3d45e955_web.html"
DEEP_LAW_SOURCE = "https://www.12377.cn/bmgz/2022/7f6f50df_web.html"

DEEP_ANNOUNCEMENTS = [
    ("2023年6月", "2023-06-20", "https://www.cac.gov.cn/2023-06/20/c_1688910683316256.htm"),
    ("2023年8月", "2023-09-01", "https://www.cac.gov.cn/2023-09/01/c_1695224377544009.htm"),
    ("2024年1月", "2024-01-05", "https://www.cac.gov.cn/2024-01/05/c_1706119043746644.htm"),
    ("2024年2月", "2024-02-18", "https://www.cac.gov.cn/2024-02/18/c_1709925427424332.htm"),
    ("2024年4月", "2024-04-11", "https://www.cac.gov.cn/2024-04/11/c_1714509267496697.htm"),
    ("2024年6月", "2024-06-12", "https://www.cac.gov.cn/2024-06/12/c_1719783421546747.htm"),
    ("2024年8月", "2024-08-05", "https://www.cac.gov.cn/2024-08/05/c_1724541639039621.htm"),
    ("2024年10月", "2024-11-01", "https://www.cac.gov.cn/2024-11/01/c_1732152604917193.htm"),
    ("2024年12月", "2024-12-20", "https://www.cac.gov.cn/2024-12/20/c_1736389545949567.htm"),
    ("2025年3月", "2025-03-12", "https://www.cac.gov.cn/2025-03/12/c_1743480314931271.htm"),
    ("2025年5月", "2025-05-19", "https://www.cac.gov.cn/2025-05/19/c_1749365589879703.htm"),
    ("2025年7月", "2025-07-14", "https://www.cac.gov.cn/2025-07/14/c_1754207718303963.htm"),
    ("2025年9月", "2025-09-11", "https://www.cac.gov.cn/2025-09/11/c_1759222331638208.htm"),
    ("2025年11月", "2025-11-06", "https://www.cac.gov.cn/2025-11/06/c_1764156698314535.htm"),
    ("2026年1月", "2026-01-07", "https://www.cac.gov.cn/2026-01/07/c_1769516642440314.htm"),
    ("2026年3月", "2026-03-12", "https://www.cac.gov.cn/2026-03/12/c_1774965172703810.htm"),
    ("2026年5月", "2026-05-06", "https://www.cac.gov.cn/2026-05/06/c_1779809434590762.htm"),
]

PROVINCE_BY_PREFIX = {
    "11": "北京", "12": "天津", "13": "河北", "14": "山西", "15": "内蒙古",
    "21": "辽宁", "22": "吉林", "23": "黑龙江", "31": "上海", "32": "江苏",
    "33": "浙江", "34": "安徽", "35": "福建", "36": "江西", "37": "山东",
    "41": "河南", "42": "湖北", "43": "湖南", "44": "广东", "45": "广西",
    "46": "海南", "50": "重庆", "51": "四川", "52": "贵州", "53": "云南",
    "54": "西藏", "61": "陕西", "62": "甘肃", "63": "青海", "64": "宁夏",
    "65": "新疆",
}

DOMAIN_RULES = [
    ("新闻资讯", ["新闻", "资讯", "文章", "信息流", "热点", "时政"]),
    ("短视频/直播", ["短视频", "视频", "直播", "视听", "音视频", "频道"]),
    ("电商/本地生活", ["电商", "购物", "商品", "营销", "外卖", "团购", "商家", "消费"]),
    ("社交/社区", ["社交", "社区", "好友", "话题", "评论", "互动", "论坛"]),
    ("搜索/浏览器", ["搜索", "检索", "浏览器", "问答"]),
    ("出行/交通", ["出行", "打车", "导航", "交通", "车辆", "网约车", "地图"]),
    ("招聘/人力资源", ["招聘", "岗位", "简历", "求职", "人才"]),
    ("教育/学习", ["教育", "学习", "课程", "题库", "作业", "培训"]),
    ("医疗/健康", ["医疗", "健康", "问诊", "医生", "药", "医院"]),
    ("金融", ["金融", "银行", "保险", "证券", "信贷", "理财"]),
    ("办公/生产力", ["办公", "文档", "写作", "会议", "协同", "生产力"]),
    ("生成合成", ["生成", "合成", "AIGC", "大模型", "对话", "绘图", "数字人", "语音", "人脸"]),
    ("安全/风控", ["风控", "安全", "反诈", "审核", "识别", "检测"]),
    ("游戏/娱乐", ["游戏", "娱乐", "音乐", "动漫", "小说"]),
]

EXPECTED_DOMESTIC = ["序号", "算法名称", "算法类别", "主体名称", "应用产品", "主要用途", "备案编号", "备注"]
EXPECTED_DEEP = ["序号", "算法名称", "角色", "主体名称", "应用产品", "主要用途", "备案编号"]


@dataclass
class Source:
    title: str
    batch: str
    date: str
    url: str
    localFile: str | None = None
    status: str = "loaded"
    message: str = ""


def clean_text(value: str | None) -> str:
    if not value:
        return ""
    return re.sub(r"\s+", " ", value.replace("\u3000", " ")).strip()


def read_docx_table(path: Path) -> list[list[str]]:
    doc = Document(str(path))
    if not doc.tables:
        raise ValueError(f"{path.name} has no table")
    return [[clean_text(cell.text) for cell in row.cells] for row in doc.tables[0].rows]


def batch_from_name(path: Path) -> str:
    match = re.search(r"（(\d{4}年\d{1,2}月)）", path.name)
    if not match:
        raise ValueError(f"Cannot read batch from file name: {path.name}")
    return match.group(1)


def batch_date(batch: str) -> str:
    match = re.match(r"(\d{4})年(\d{1,2})月", batch)
    return f"{match.group(1)}-{int(match.group(2)):02d}" if match else batch


def province_from_record_number(record_number: str) -> str:
    match = re.search(r"网信算备(\d{6})", record_number)
    return PROVINCE_BY_PREFIX.get(match.group(1)[:2], "未识别") if match else "未识别"


def domain_tags(*parts: str) -> list[str]:
    text = " ".join(parts).lower()
    tags = [label for label, words in DOMAIN_RULES if any(word.lower() in text for word in words)]
    return tags or ["未分类/待复核"]


def stable_id(*parts: str) -> str:
    raw = "|".join(clean_text(part) for part in parts)
    return hashlib.sha1(raw.encode("utf-8")).hexdigest()[:16]


def normalize_domestic(path: Path) -> tuple[list[dict], Source]:
    rows = read_docx_table(path)
    if rows[0][: len(EXPECTED_DOMESTIC)] != EXPECTED_DOMESTIC:
        raise ValueError(f"{path.name} header mismatch: {rows[0]}")
    batch = batch_from_name(path)
    records = []
    for row in rows[1:]:
        cells = (row + [""] * len(EXPECTED_DOMESTIC))[: len(EXPECTED_DOMESTIC)]
        if not cells[0] or not cells[1]:
            continue
        record_number = cells[6]
        records.append({
            "id": stable_id("domestic", record_number, cells[1], cells[3]),
            "filingType": "互联网信息服务算法备案",
            "batch": batch,
            "announcementDate": batch_date(batch),
            "algorithmName": cells[1],
            "algorithmClass": cells[2],
            "role": "",
            "entityName": cells[3],
            "product": cells[4],
            "purpose": cells[5],
            "recordNumber": record_number,
            "remark": cells[7],
            "province": province_from_record_number(record_number),
            "domainTags": domain_tags(cells[4], cells[5], cells[2]),
            "sourceUrl": DOMESTIC_SOURCE,
            "sourceFile": path.name,
        })
    return records, Source("境内互联网信息服务算法备案清单", batch, batch_date(batch), DOMESTIC_SOURCE, path.name)


def normalize_deep(path: Path, fallback_batch: str | None = None, fallback_date: str | None = None, source_url: str = "") -> tuple[list[dict], Source]:
    rows = read_docx_table(path)
    if len(rows[0]) >= 9 and rows[0][5] == "主要用途" and rows[0][6] == "主要用途":
        rows = [row[:6] + row[7:] for row in rows]
    if rows[0][: len(EXPECTED_DEEP)] != EXPECTED_DEEP:
        raise ValueError(f"{path.name} header mismatch: {rows[0]}")
    batch = fallback_batch or batch_from_name(path)
    records = []
    for row in rows[1:]:
        cells = (row + [""] * 8)[:8]
        if not cells[1]:
            continue
        record_number = cells[6]
        records.append({
            "id": stable_id("deep", record_number, cells[1], cells[3], cells[2]),
            "filingType": "深度合成服务算法备案",
            "batch": batch,
            "announcementDate": fallback_date or batch_date(batch),
            "algorithmName": cells[1],
            "algorithmClass": f"深度合成-{cells[2]}",
            "role": f"深度合成-{cells[2]}",
            "entityName": cells[3],
            "product": cells[4],
            "purpose": cells[5],
            "recordNumber": record_number,
            "remark": cells[7] if len(cells) > 7 else "",
            "province": province_from_record_number(record_number),
            "domainTags": domain_tags(cells[4], cells[5], cells[1], "生成合成"),
            "sourceUrl": source_url,
            "sourceFile": path.name,
        })
    return records, Source("境内深度合成服务算法备案清单", batch, fallback_date or batch_date(batch), source_url, path.name)


def browser_session() -> requests.Session:
    session = requests.Session()
    session.headers.update({
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
    })
    return session


def fetch_html(url: str, session: requests.Session) -> str:
    last_error = ""
    for candidate in [url, url.replace("https://", "http://", 1)]:
        try:
            response = session.get(candidate, timeout=30, allow_redirects=True)
            response.encoding = response.apparent_encoding
            if response.status_code == 200 and "知道创宇云防御" not in response.text:
                return response.text
            last_error = f"{response.status_code}: {response.text[:120]}"
        except Exception as exc:
            last_error = str(exc)
        time.sleep(1)
    raise RuntimeError(f"fetch failed for {url}: {last_error}")


def discover_attachment(url: str, session: requests.Session) -> tuple[str, str]:
    soup = BeautifulSoup(fetch_html(url, session), "lxml")
    for anchor in soup.find_all("a"):
        text = clean_text(anchor.get_text(" ", strip=True))
        href = anchor.get("href")
        if href and ("备案清单" in text or href.lower().endswith((".doc", ".docx"))):
            return urljoin(url, href), text
    raise RuntimeError(f"no attachment found at {url}")


def download_one_announcement(url: str, session: requests.Session, batch: str | None = None, date: str = "") -> Source:
    attachment, label = discover_attachment(url, session)
    guessed_batch = batch
    if not guessed_batch:
        match = re.search(r"（(\d{4}年\d{1,2}月)）", label)
        guessed_batch = match.group(1) if match else time.strftime("%Y年%m月")
    target = DEEP_DIR / f"境内深度合成服务算法备案清单（{guessed_batch}）.docx"
    response = session.get(attachment, timeout=45)
    response.raise_for_status()
    if len(response.content) < 5_000:
        raise RuntimeError("attachment too small")
    target.write_bytes(response.content)
    return Source("境内深度合成服务算法备案清单", guessed_batch, date or batch_date(guessed_batch), url, target.name, "downloaded")


def download_deep_announcements(extra_url: str | None = None) -> None:
    DEEP_DIR.mkdir(parents=True, exist_ok=True)
    session = browser_session()
    if extra_url:
        try:
            download_one_announcement(extra_url, session)
        except Exception as exc:
            print(f"warning: failed to download manual announcement {extra_url}: {exc}", file=sys.stderr)
    for batch, date, url in DEEP_ANNOUNCEMENTS:
        target = DEEP_DIR / f"境内深度合成服务算法备案清单（{batch}）.docx"
        if target.exists():
            continue
        try:
            download_one_announcement(url, session, batch, date)
        except Exception as exc:
            print(f"warning: failed to download {batch} deep synthesis source {url}: {exc}", file=sys.stderr)
        time.sleep(1)


def parse_law_articles(title: str, source_url: str, effective_date: str) -> dict:
    session = browser_session()
    html = fetch_html(source_url, session)
    soup = BeautifulSoup(html, "lxml")
    lines = [clean_text(line) for line in soup.get_text("\n", strip=True).split("\n")]
    start = next((i for i, line in enumerate(lines) if line == title), 0)
    chapters = []
    current = {"chapter": "正文", "articles": []}
    article_re = re.compile(r"^(第[一二三四五六七八九十百]+条)\s*(.*)")
    for line in lines[start:]:
        if not line or line in {"分享：", "打印", "关闭", "返回顶部"}:
            continue
        if line.startswith("第") and "章" in line and len(line) < 20:
            if current["articles"]:
                chapters.append(current)
            current = {"chapter": line, "articles": []}
            continue
        article_match = article_re.match(line)
        if article_match:
            current["articles"].append({"number": article_match.group(1), "text": article_match.group(2).strip()})
        elif current["articles"]:
            current["articles"][-1]["text"] += "\n" + line
        if "本规定自" in line and "施行" in line:
            break
    if current["articles"]:
        chapters.append(current)
    if not chapters:
        raise RuntimeError(f"no law articles parsed from {source_url}")
    return {"title": title, "sourceUrl": source_url, "effectiveDate": effective_date, "chapters": chapters}


def load_supplemental_laws() -> list[dict]:
    path = OUT_DIR / "supplemental-laws.json"
    if not path.exists():
        return []
    return json.loads(path.read_text(encoding="utf-8"))


def merge_laws(base_laws: list[dict], supplemental_laws: list[dict]) -> list[dict]:
    seen = {law.get("title") for law in base_laws}
    merged = [*base_laws]
    for law in supplemental_laws:
        title = law.get("title")
        if title and title not in seen:
            merged.append(law)
            seen.add(title)
    return merged


def write_json(path: Path, payload) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_csv(path: Path, records: list[dict]) -> None:
    fields = ["filingType", "batch", "announcementDate", "algorithmName", "algorithmClass", "role", "entityName", "product", "purpose", "recordNumber", "remark", "province", "domainTags", "sourceUrl"]
    with path.open("w", encoding="utf-8-sig", newline="") as file:
        writer = csv.DictWriter(file, fieldnames=fields)
        writer.writeheader()
        for record in records:
            row = {field: record.get(field, "") for field in fields}
            row["domainTags"] = "、".join(record.get("domainTags", []))
            writer.writerow(row)


def build(fetch_remote: bool, announcement_url: str | None = None) -> None:
    if fetch_remote:
        download_deep_announcements(announcement_url)

    records: list[dict] = []
    sources: list[Source] = []
    domestic_files = sorted(DOMESTIC_DIR.glob("*.docx"), key=lambda p: (batch_date(batch_from_name(p)), p.name))
    if len(domestic_files) != 18:
        raise RuntimeError(f"Expected 18 domestic docx files, found {len(domestic_files)}")
    for file in domestic_files:
        parsed, source = normalize_domestic(file)
        records.extend(parsed)
        sources.append(source)

    announcement_by_batch = {batch: (date, url) for batch, date, url in DEEP_ANNOUNCEMENTS}
    for file in sorted(DEEP_DIR.glob("*.docx"), key=lambda p: (batch_date(batch_from_name(p)), p.name)):
        batch = batch_from_name(file)
        date, url = announcement_by_batch.get(batch, (batch_date(batch), ""))
        parsed, source = normalize_deep(file, batch, date, url)
        records.extend(parsed)
        sources.append(source)

    unique_records = []
    seen = {}
    duplicate_count = 0
    for record in records:
        key = (record["filingType"], record["recordNumber"], record["algorithmName"], record["entityName"])
        if key in seen:
            duplicate_count += 1
            previous = seen[key]
            previous.setdefault("duplicateBatches", [previous["batch"]])
            if record["batch"] not in previous["duplicateBatches"]:
                previous["duplicateBatches"].append(record["batch"])
            continue
        seen[key] = record
        unique_records.append(record)
    records = unique_records

    existing_laws_path = OUT_DIR / "laws.json"
    try:
        laws = [
            parse_law_articles("互联网信息服务算法推荐管理规定", ALGO_LAW_SOURCE, "2022-03-01"),
            parse_law_articles("互联网信息服务深度合成管理规定", DEEP_LAW_SOURCE, "2023-01-10"),
        ]
    except Exception as exc:
        if not existing_laws_path.exists():
            raise
        print(f"warning: law source fetch failed, reusing existing laws.json: {exc}", file=sys.stderr)
        laws = json.loads(existing_laws_path.read_text(encoding="utf-8"))
    laws = merge_laws(laws, load_supplemental_laws())

    facets = {
        "filingTypes": sorted({r["filingType"] for r in records}),
        "batches": sorted({r["batch"] for r in records}, key=batch_date),
        "algorithmClasses": sorted({r["algorithmClass"] for r in records if r["algorithmClass"]}),
        "provinces": sorted({r["province"] for r in records if r["province"]}),
        "domainTags": sorted({tag for r in records for tag in r["domainTags"]}),
    }
    stats = {
        "recordCount": len(records),
        "domesticCount": sum(1 for r in records if r["filingType"] == "互联网信息服务算法备案"),
        "deepCount": sum(1 for r in records if r["filingType"] == "深度合成服务算法备案"),
        "sourceCount": len(sources),
        "duplicateCount": duplicate_count,
        "generatedAt": time.strftime("%Y-%m-%d %H:%M:%S"),
    }

    write_json(OUT_DIR / "records.json", records)
    write_json(OUT_DIR / "sources.json", [source.__dict__ for source in sources])
    write_json(OUT_DIR / "laws.json", laws)
    write_json(OUT_DIR / "facets.json", facets)
    write_json(OUT_DIR / "stats.json", stats)
    write_csv(OUT_DIR / "records.csv", records)
    print(json.dumps(stats, ensure_ascii=False))


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--fetch-remote", action="store_true")
    parser.add_argument("--announcement-url")
    args = parser.parse_args()
    try:
        build(args.fetch_remote or bool(args.announcement_url), args.announcement_url)
    except Exception as exc:
        print(f"build_data failed: {exc}", file=sys.stderr)
        raise
